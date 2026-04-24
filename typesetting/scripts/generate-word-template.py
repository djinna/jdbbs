#!/usr/bin/env python3
"""Generate a styled Word (.docx) template from a book spec JSON.

Usage:
    echo '{...}' | python3 generate-word-template.py > template.docx
    python3 generate-word-template.py --spec-file spec.json --output template.docx

Reads a book_specs.data JSON on stdin (or --spec-file) and produces a .docx
with paragraph/character styles matching the spec. The .docx serves as a
starting template for copyeditors working in Word.
"""

import argparse
import io
import json
import sys

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Emu
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

SECTION_BREAK_CHARS = {
    "breve":    "\u02D8",       # ˘
    "asterism":  "\u2042",      # ⁂
    "dinkus":   "* * *",
    "fleuron":  "\u2767",       # ❧
    "rule":     "\u2014" * 3,   # ———
}


def _weight_to_bold(weight) -> bool | None:
    """Convert a heading weight spec to a python-docx bold flag.

    Returns True for bold/semibold/600, False for medium/500/normal, None for
    anything else (inherit).
    """
    if weight is None:
        return None
    w = str(weight).strip().lower()
    if w in ("bold", "700"):
        return True
    if w in ("semibold", "600"):
        return True
    if w in ("medium", "500", "normal", "400", "regular"):
        return False
    # numeric fallback
    try:
        return int(w) >= 600
    except ValueError:
        return None


def _ensure_font_exists(run_font, font_name: str):
    """Set the font name on a run's font object (western + eastAsia)."""
    run_font.name = font_name
    # Also set the eastAsia font so Word doesn't substitute
    r = run_font.element
    r.set(qn("w:eastAsia"), font_name)


def _set_paragraph_style_font(style, font_name: str, size_pt: float,
                                bold: bool | None = None,
                                italic: bool | None = None):
    """Configure the font for a paragraph style."""
    font = style.font
    _ensure_font_exists(font, font_name)
    font.size = Pt(size_pt)
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic


def _set_line_spacing_pt(style, spacing_pt: float):
    """Set exact line spacing in points on a paragraph style."""
    pf = style.paragraph_format
    pf.line_spacing = Pt(spacing_pt)
    # python-docx sets line_spacing_rule automatically when given Pt


def _set_first_line_indent(style, indent_pt: float):
    """Set first-line indent in points."""
    style.paragraph_format.first_line_indent = Pt(indent_pt)


def _set_left_indent(style, indent_pt: float):
    """Set left indent in points."""
    style.paragraph_format.left_indent = Pt(indent_pt)


def _set_alignment(style, justify: bool):
    """Set alignment to JUSTIFY or LEFT."""
    style.paragraph_format.alignment = (
        WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    )


# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------

def build_template(spec: dict) -> Document:
    """Build and return a Document configured from *spec*."""
    doc = Document()

    meta = spec.get("metadata", {})
    typo = spec.get("typography", {})
    hdgs = spec.get("headings", {})
    elms = spec.get("elements", {})
    customs = spec.get("custom_styles", [])

    # Derived values
    body_font    = typo.get("body_font", "Libertinus Serif")
    heading_font = typo.get("heading_font", "Source Sans 3")
    code_font    = typo.get("code_font", "JetBrains Mono")
    base_size    = float(typo.get("base_size_pt", 10))
    leading      = float(typo.get("leading_pt", 2))
    line_sp      = base_size + leading
    indent_em    = float(typo.get("paragraph_indent_em", 0.75))
    indent_pt    = indent_em * base_size
    justify      = bool(typo.get("justify", True))

    # ------------------------------------------------------------------
    # 1. Normal (base body text) style
    # ------------------------------------------------------------------
    normal = doc.styles["Normal"]
    _set_paragraph_style_font(normal, body_font, base_size)
    _set_line_spacing_pt(normal, line_sp)
    _set_first_line_indent(normal, indent_pt)
    _set_alignment(normal, justify)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    # ------------------------------------------------------------------
    # 2. First Paragraph (no indent, used after headings / breaks)
    # ------------------------------------------------------------------
    first_para = doc.styles.add_style("First Paragraph", WD_STYLE_TYPE.PARAGRAPH)
    first_para.base_style = normal
    _set_first_line_indent(first_para, 0)

    # ------------------------------------------------------------------
    # 3. Heading styles
    # ------------------------------------------------------------------
    heading_specs = [
        ("Heading 1", hdgs.get("h1_size_em", 1.667), hdgs.get("h1_weight", "bold")),
        ("Heading 2", hdgs.get("h2_size_em", 1.333), hdgs.get("h2_weight", 600)),
        ("Heading 3", hdgs.get("h3_size_em", 1.0),   hdgs.get("h3_weight", "medium")),
    ]
    for style_name, size_em, weight in heading_specs:
        style = doc.styles[style_name]
        size_pt = float(size_em) * base_size
        bold = _weight_to_bold(weight)
        _set_paragraph_style_font(style, heading_font, size_pt, bold=bold)
        style.paragraph_format.space_before = Pt(base_size * 1.5)
        style.paragraph_format.space_after = Pt(base_size * 0.5)
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Keep with next
        style.paragraph_format.keep_with_next = True

    # ------------------------------------------------------------------
    # 4. Block Quote
    # ------------------------------------------------------------------
    bq_style = doc.styles.add_style("Block Quote", WD_STYLE_TYPE.PARAGRAPH)
    bq_style.base_style = normal
    bq_italic = elms.get("blockquote_style", "italic") == "italic"
    _set_paragraph_style_font(bq_style, body_font, base_size, italic=bq_italic)
    _set_left_indent(bq_style, base_size * 2)  # ~2em left indent
    _set_first_line_indent(bq_style, 0)
    bq_style.paragraph_format.space_before = Pt(base_size * 0.5)
    bq_style.paragraph_format.space_after = Pt(base_size * 0.5)

    # ------------------------------------------------------------------
    # 5. Code Block
    # ------------------------------------------------------------------
    code_size_em = float(elms.get("code_block_size_em", 0.8))
    code_size = code_size_em * base_size

    cb_style = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    cb_style.base_style = normal
    _set_paragraph_style_font(cb_style, code_font, code_size)
    _set_left_indent(cb_style, base_size * 1.5)
    _set_first_line_indent(cb_style, 0)
    _set_alignment(cb_style, False)  # left-aligned
    _set_line_spacing_pt(cb_style, code_size + leading)
    cb_style.paragraph_format.space_before = Pt(base_size * 0.5)
    cb_style.paragraph_format.space_after = Pt(base_size * 0.5)

    # ------------------------------------------------------------------
    # 6. Section Break
    # ------------------------------------------------------------------
    sb_char_key = elms.get("section_break", "breve")
    sb_char = SECTION_BREAK_CHARS.get(sb_char_key, sb_char_key)

    sb_style = doc.styles.add_style("Section Break", WD_STYLE_TYPE.PARAGRAPH)
    sb_style.base_style = normal
    sb_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_first_line_indent(sb_style, 0)
    sb_style.paragraph_format.space_before = Pt(base_size)
    sb_style.paragraph_format.space_after = Pt(base_size)

    # ------------------------------------------------------------------
    # 7. Poem / Verse
    # ------------------------------------------------------------------
    poem_size_em = float(elms.get("poem_size_em", 0.75))
    poem_size = poem_size_em * base_size

    poem_style = doc.styles.add_style("Verse", WD_STYLE_TYPE.PARAGRAPH)
    poem_style.base_style = normal
    _set_paragraph_style_font(poem_style, body_font, poem_size, italic=True)
    _set_left_indent(poem_style, base_size * 2)
    _set_first_line_indent(poem_style, 0)
    _set_alignment(poem_style, False)
    _set_line_spacing_pt(poem_style, poem_size + leading)
    poem_style.paragraph_format.space_before = Pt(base_size * 0.25)
    poem_style.paragraph_format.space_after = Pt(base_size * 0.25)

    # ------------------------------------------------------------------
    # 8. Copyright
    # ------------------------------------------------------------------
    cr_style = doc.styles.add_style("Copyright", WD_STYLE_TYPE.PARAGRAPH)
    cr_style.base_style = normal
    cr_size = base_size * 0.8
    _set_paragraph_style_font(cr_style, body_font, cr_size)
    _set_first_line_indent(cr_style, 0)
    _set_alignment(cr_style, False)
    _set_line_spacing_pt(cr_style, cr_size + leading)

    # ------------------------------------------------------------------
    # 9. Epigraph
    # ------------------------------------------------------------------
    ep_style = doc.styles.add_style("Epigraph", WD_STYLE_TYPE.PARAGRAPH)
    ep_style.base_style = normal
    _set_paragraph_style_font(ep_style, body_font, base_size, italic=True)
    _set_left_indent(ep_style, base_size * 3)
    _set_first_line_indent(ep_style, 0)
    _set_alignment(ep_style, False)
    ep_style.paragraph_format.space_before = Pt(base_size * 0.5)
    ep_style.paragraph_format.space_after = Pt(base_size * 0.5)

    # ------------------------------------------------------------------
    # 10. Custom styles from spec
    # ------------------------------------------------------------------
    for cs in customs:
        cs_name = cs.get("word_style") or cs.get("name", "Custom")
        cs_type_str = cs.get("type", "paragraph").lower()
        # Skip if this style already exists (e.g. built-in or already created)
        if cs_name in [s.name for s in doc.styles]:
            continue
        if cs_type_str == "character":
            cs_style = doc.styles.add_style(cs_name, WD_STYLE_TYPE.CHARACTER)
            cs_style.base_style = doc.styles["Default Paragraph Font"]
        else:
            cs_style = doc.styles.add_style(cs_name, WD_STYLE_TYPE.PARAGRAPH)
            cs_style.base_style = normal
        # Set description if present (stored as style's element name hint)
        # python-docx doesn't expose style description directly, but we note
        # it in the sample content below.

    # ------------------------------------------------------------------
    # Sample content demonstrating each style
    # ------------------------------------------------------------------
    title = meta.get("title", "Book Title")
    author = meta.get("author", "Author Name")

    # Title page
    p = doc.add_heading(title, level=1)
    p = doc.add_paragraph(f"by {author}", style="First Paragraph")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # blank line

    # Instructions
    p = doc.add_heading("Template Guide", level=2)

    p = doc.add_paragraph(
        "This document is a styled Word template generated from your book "
        "specification. Each paragraph style below matches a Typst style used "
        "in final production. Use these styles consistently so the manuscript "
        "converts cleanly.",
        style="First Paragraph"
    )

    # --- Normal / Body Text ---
    p = doc.add_heading("Body Text (Normal)", level=3)
    doc.add_paragraph(
        "This paragraph uses the Normal style — the default for body text. "
        f"Font: {body_font} at {base_size}pt with {line_sp}pt line spacing. "
        f"First-line indent: {indent_pt:.1f}pt. "
        f"{'Justified' if justify else 'Left-aligned'}.",
        style="First Paragraph"
    )
    doc.add_paragraph(
        "Subsequent paragraphs in the same section use Normal style with the "
        "first-line indent. Only the first paragraph after a heading or break "
        "should use 'First Paragraph' (no indent)."
    )

    # --- First Paragraph ---
    p = doc.add_heading("First Paragraph", level=3)
    doc.add_paragraph(
        "Apply 'First Paragraph' to the first paragraph after any heading or "
        "section break. It is identical to Normal but has no first-line indent.",
        style="First Paragraph"
    )

    # --- Block Quote ---
    p = doc.add_heading("Block Quote", level=3)
    doc.add_paragraph(
        "Use 'Block Quote' for extended quotations:",
        style="First Paragraph"
    )
    doc.add_paragraph(
        "This is a sample block quotation. It is indented from the left margin "
        "and set in italic (per your spec). Use this style for any quotation "
        "that runs longer than a few words.",
        style="Block Quote"
    )

    # --- Code Block ---
    p = doc.add_heading("Code Block", level=3)
    doc.add_paragraph(
        "Use 'Code Block' for terminal output, code listings, or monospaced content:",
        style="First Paragraph"
    )
    doc.add_paragraph(
        f"$ echo 'Hello, World!'\n"
        f"Hello, World!\n"
        f"# Font: {code_font} at {code_size:.1f}pt",
        style="Code Block"
    )

    # --- Section Break ---
    p = doc.add_heading("Section Break", level=3)
    doc.add_paragraph(
        "Use 'Section Break' to mark scene or section divisions. "
        f"Your spec uses the '{sb_char_key}' symbol:",
        style="First Paragraph"
    )
    doc.add_paragraph(sb_char, style="Section Break")
    doc.add_paragraph(
        "The paragraph after a section break should use 'First Paragraph'.",
        style="First Paragraph"
    )

    # --- Verse ---
    p = doc.add_heading("Verse / Poem", level=3)
    doc.add_paragraph(
        "Use 'Verse' for poetry or lyrics. Each line is a separate line within "
        "one paragraph (Shift+Enter for soft returns):",
        style="First Paragraph"
    )
    doc.add_paragraph(
        "Shall I compare thee to a summer's day?\n"
        "Thou art more lovely and more temperate.",
        style="Verse"
    )

    # --- Epigraph ---
    p = doc.add_heading("Epigraph", level=3)
    doc.add_paragraph(
        "Use 'Epigraph' for chapter-opening quotations:",
        style="First Paragraph"
    )
    doc.add_paragraph(
        "In the beginning was the Word.\n— John 1:1",
        style="Epigraph"
    )

    # --- Copyright ---
    p = doc.add_heading("Copyright", level=3)
    doc.add_paragraph(
        "Use 'Copyright' for the copyright page:",
        style="First Paragraph"
    )
    publisher = meta.get("publisher", "Publisher")
    doc.add_paragraph(
        f"Copyright © 2025 {author}. All rights reserved.\n"
        f"Published by {publisher}.\n"
        "No part of this book may be reproduced without permission.",
        style="Copyright"
    )

    # --- Custom Styles ---
    if customs:
        p = doc.add_heading("Custom Styles", level=2)
        for cs in customs:
            cs_name = cs.get("word_style") or cs.get("name", "Custom")
            desc = cs.get("description", "No description provided.")
            cs_type_str = cs.get("type", "paragraph").lower()
            if cs_type_str == "character":
                p = doc.add_paragraph(style="First Paragraph")
                p.add_run(f"{cs_name} (character style): ").bold = True
                # Try to apply the character style to a sample run
                try:
                    r = p.add_run("sample text")
                    r.style = doc.styles[cs_name]
                except KeyError:
                    p.add_run("sample text")
                p.add_run(f" — {desc}")
            else:
                p = doc.add_paragraph(style="First Paragraph")
                p.add_run(f"{cs_name}: ").bold = True
                p.add_run(desc)
                # Add a sample paragraph in the style
                try:
                    doc.add_paragraph(
                        f"This paragraph uses the '{cs_name}' style.",
                        style=cs_name
                    )
                except KeyError:
                    pass

    # --- Style summary table ---
    p = doc.add_heading("Style Summary", level=2)

    style_rows = [
        ("Normal",          f"{body_font}, {base_size}pt",    "Default body text with indent"),
        ("First Paragraph", f"{body_font}, {base_size}pt",    "After headings/breaks (no indent)"),
        ("Heading 1",       f"{heading_font}, {float(hdgs.get('h1_size_em', 1.667)) * base_size:.1f}pt", "Chapter titles"),
        ("Heading 2",       f"{heading_font}, {float(hdgs.get('h2_size_em', 1.333)) * base_size:.1f}pt", "Sub-sections"),
        ("Heading 3",       f"{heading_font}, {float(hdgs.get('h3_size_em', 1.0)) * base_size:.1f}pt",   "Sub-sub-sections"),
        ("Block Quote",     f"{body_font}, {base_size}pt italic", "Extended quotations"),
        ("Code Block",      f"{code_font}, {code_size:.1f}pt",    "Code / terminal output"),
        ("Section Break",   f"Centered, '{sb_char_key}'",         "Scene / section divider"),
        ("Verse",           f"{body_font}, {poem_size:.1f}pt italic", "Poetry / lyrics"),
        ("Epigraph",        f"{body_font}, {base_size}pt italic", "Chapter-opening quotation"),
        ("Copyright",       f"{body_font}, {cr_size:.1f}pt",      "Copyright page text"),
    ]
    for cs in customs:
        cs_name = cs.get("word_style") or cs.get("name", "Custom")
        cs_desc = cs.get("description", "")
        style_rows.append((cs_name, cs.get("type", "paragraph"), cs_desc))

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Style Name"
    hdr[1].text = "Font / Size"
    hdr[2].text = "Usage"
    for name, font_info, usage in style_rows:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = font_info
        row[2].text = usage

    return doc


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Generate a Word template from a book spec JSON."
    )
    parser.add_argument(
        "--spec-file", "-s",
        help="Path to spec JSON file (reads stdin if omitted)",
    )
    parser.add_argument(
        "--output", "-o",
        help="Output .docx path (writes to stdout if omitted)",
    )
    args = parser.parse_args()

    # Read spec
    if args.spec_file:
        with open(args.spec_file, "r", encoding="utf-8") as f:
            spec = json.load(f)
    else:
        spec = json.load(sys.stdin)

    # Build document
    doc = build_template(spec)

    # Write output
    if args.output:
        doc.save(args.output)
    else:
        buf = io.BytesIO()
        doc.save(buf)
        sys.stdout.buffer.write(buf.getvalue())


if __name__ == "__main__":
    main()
