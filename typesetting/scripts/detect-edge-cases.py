#!/usr/bin/env python3
"""
Edge Case Detection and Review System
Detects manual formatting in Word documents and presents them for review
"""

import argparse
import json
import re
import unicodedata
from pathlib import Path
from typing import List, Dict, Tuple
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor
from datetime import datetime
import html


class EdgeCaseDetector:
    """Detects manual/local formatting that might need review"""
    
    ASCII_DENSE_RE = re.compile(r"[\\/_|()\[\]{}<>-]{3,}")
    CJK_RE = re.compile(r"[\u4E00-\u9FFF\u3040-\u30FF\u3400-\u4DBF]")
    THAI_RE = re.compile(r"[\u0E00-\u0E7F]")
    BUILTIN_PARAGRAPH_STYLES = {
        'normal', 'body text', 'first paragraph', 'heading 1', 'heading 2', 'heading 3',
        'title', 'subtitle', 'quote', 'block quote', 'list paragraph'
    }
    BUILTIN_CHARACTER_STYLES = {
        'default paragraph font', 'strong', 'emphasis', 'subtle emphasis', 'intense emphasis'
    }

    def __init__(self, doc_path: str, declared_styles: List[Dict] | None = None):
        self.doc = Document(doc_path)
        self.edge_cases = []
        self.doc_path = doc_path
        self.declared_styles = declared_styles or []
        self.declared_style_names = {
            self._normalize_style_name((item.get('word_style') or item.get('name') or ''))
            for item in self.declared_styles
            if self._normalize_style_name((item.get('word_style') or item.get('name') or ''))
        }
        self.template_guide_cutoff = self._find_template_guide_cutoff()

    def _paragraph_context(self, index: int) -> Dict:
        context_before = []
        context_after = []
        for offset in (-1, -2):
            pos = index + offset
            if 0 <= pos < len(self.doc.paragraphs):
                text = (self.doc.paragraphs[pos].text or '').strip()
                if text:
                    context_before.append(text[:120])
        context_before.reverse()
        for offset in (1, 2):
            pos = index + offset
            if 0 <= pos < len(self.doc.paragraphs):
                text = (self.doc.paragraphs[pos].text or '').strip()
                if text:
                    context_after.append(text[:120])
        return {
            'context_before': context_before,
            'context_after': context_after,
        }
        
    def detect_all(self) -> List[Dict]:
        """Run all detection methods"""
        self.detect_manual_formatting()
        self.detect_unusual_fonts()
        self.detect_colored_text()
        self.detect_manual_lists()
        self.detect_manual_breaks()
        self.detect_mixed_styles()
        self.detect_direct_formatting()
        self.detect_image_inventory()
        self.detect_observed_styles()
        self.detect_special_typography()
        self.detect_language_scripts()
        return self.edge_cases
    
    def detect_manual_formatting(self):
        """Detect bold/italic applied directly (not via character style)"""
        for i, para in enumerate(self.doc.paragraphs):
            if self._should_skip_paragraph(para):
                continue
            for run in para.runs:
                issues = []
                
                # Check for direct bold (not from style)
                if run.bold and not self._is_from_style(run, 'bold'):
                    issues.append("manual bold")
                    
                # Check for direct italic (not from style)  
                if run.italic and not self._is_from_style(run, 'italic'):
                    issues.append("manual italic")
                    
                # Check for direct underline
                if run.underline:
                    issues.append("manual underline")
                    
                if issues:
                    self.edge_cases.append({
                        'type': 'manual_formatting',
                        'location': f'Paragraph {i+1}',
                        'text': run.text[:100],
                        'issues': issues,
                        'severity': 'medium',
                        'auto_decision': 'preserve',
                        'suggestion': 'Manual emphasis detected — preserve it as intentional bold/italic formatting in EPUB and Typst'
                    })
    
    def detect_unusual_fonts(self):
        """Detect non-standard fonts, but aggregate emoji-heavy cases instead of spamming runs."""
        standard_fonts = {
            'Calibri', 'Times New Roman', 'Arial', 'Libertinus Serif',
            'Source Sans 3', 'JetBrains Mono', 'Cambria', 'Georgia'
        }
        unusual_runs = []
        for i, para in enumerate(self.doc.paragraphs):
            if self._should_skip_paragraph(para):
                continue
            for run in para.runs:
                if run.font.name and run.font.name not in standard_fonts:
                    unusual_runs.append((i + 1, run))

        grouped = {}
        for location, run in unusual_runs:
            font_name = run.font.name
            text = (run.text or '').strip()
            key = font_name
            grouped.setdefault(key, []).append((location, text))

        for font_name, entries in grouped.items():
            nonempty_samples = [text for _, text in entries if text]
            locations = sorted({loc for loc, _ in entries})
            if self._font_should_be_aggregated(font_name, nonempty_samples):
                self.edge_cases.append({
                    'type': 'font_treatment',
                    'subtype': 'emoji_font_usage',
                    'location': f'Paragraphs {locations[0]}-{locations[-1]}' if len(locations) > 1 else f'Paragraph {locations[0]}',
                    'text': '; '.join(nonempty_samples[:5]) if nonempty_samples else font_name,
                    'font': font_name,
                    'severity': 'low',
                    'suggestion': 'Emoji/special font usage detected — absorb as character-style treatment and verify font survival in EPUB and PDF'
                })
                continue
            for location, text in entries:
                self.edge_cases.append({
                    'type': 'unusual_font',
                    'location': f'Paragraph {location}',
                    'text': text[:100] if text else font_name,
                    'font': font_name,
                    'severity': 'high',
                    'suggestion': f'Non-standard font "{font_name}" - review for conversion'
                })
    
    def detect_colored_text(self):
        """Detect colored text (non-black)"""
        for i, para in enumerate(self.doc.paragraphs):
            if self._should_skip_paragraph(para):
                continue
            for run in para.runs:
                # Check font color
                if run.font.color and run.font.color.rgb:
                    rgb = run.font.color.rgb
                    if rgb != RGBColor(0, 0, 0):  # Not black
                        self.edge_cases.append({
                            'type': 'colored_text',
                            'location': f'Paragraph {i+1}',
                            'text': run.text[:100],
                            'color': f'RGB({rgb[0]},{rgb[1]},{rgb[2]})',
                            'severity': 'high',
                            'suggestion': 'Colored text detected - will be removed for print'
                        })
                
                # Check highlight color
                if run.font.highlight_color and run.font.highlight_color != WD_COLOR_INDEX.AUTO:
                    self.edge_cases.append({
                        'type': 'highlighted_text',
                        'location': f'Paragraph {i+1}',
                        'text': run.text[:100],
                        'severity': 'medium',
                        'suggestion': 'Highlighted text detected - review purpose'
                    })
    
    def detect_manual_lists(self):
        """Detect manually formatted lists (using -, *, •, numbers)"""
        list_patterns = [
            (r'^\s*[-*•]\s+', 'bullet'),
            (r'^\s*\d+[\.)]\s+', 'numbered'),
            (r'^\s*[a-zA-Z][\.)]\s+', 'lettered')
        ]
        
        import re
        for i, para in enumerate(self.doc.paragraphs):
            if self._should_skip_paragraph(para):
                continue
            text = para.text.strip()
            for pattern, list_type in list_patterns:
                if re.match(pattern, text):
                    self.edge_cases.append({
                        'type': 'manual_list',
                        'location': f'Paragraph {i+1}',
                        'text': text[:100],
                        'list_type': list_type,
                        'severity': 'low',
                        'suggestion': f'Manual {list_type} list - convert to proper list style'
                    })
                    break
    
    def detect_manual_breaks(self):
        """Detect manual section breaks (*, * * *, ---, ___, etc.)"""
        break_patterns = [
            r'^\s*\*\s*\*\s*\*\s*$',
            r'^\s*[-_]{3,}\s*$',
            r'^\s*[•·∗]{3,}\s*$',
            r'^\s*\*\s*$',
            r'^#{3,}\s*$'
        ]
        
        import re
        for i, para in enumerate(self.doc.paragraphs):
            if self._should_skip_paragraph(para):
                continue
            text = para.text.strip()
            for pattern in break_patterns:
                if re.match(pattern, text):
                    self.edge_cases.append({
                        'type': 'manual_break',
                        'location': f'Paragraph {i+1}',
                        'text': text,
                        'severity': 'low',
                        'suggestion': 'Manual section break - convert to proper style or spacing'
                    })
                    break
    
    def detect_mixed_styles(self):
        """Detect paragraphs with multiple fonts/sizes"""
        for i, para in enumerate(self.doc.paragraphs):
            if self._should_skip_paragraph(para):
                continue
            fonts = set()
            sizes = set()
            
            for run in para.runs:
                if run.font.name:
                    fonts.add(run.font.name)
                if run.font.size:
                    sizes.add(run.font.size)
            
            if len(fonts) > 1 or len(sizes) > 1:
                self.edge_cases.append({
                    'type': 'mixed_formatting',
                    'location': f'Paragraph {i+1}',
                    'text': para.text[:100],
                    'fonts': list(fonts),
                    'sizes': [str(s) for s in sizes],
                    'severity': 'medium',
                    'suggestion': 'Multiple fonts/sizes in one paragraph - review intent'
                })
    
    def detect_direct_formatting(self):
        """Detect spacing, indentation, alignment issues"""
        for i, para in enumerate(self.doc.paragraphs):
            if self._should_skip_paragraph(para):
                continue
            issues = []
            
            # Check for manual spacing
            if para.paragraph_format.space_before:
                issues.append(f"space before: {para.paragraph_format.space_before}")
            if para.paragraph_format.space_after:
                issues.append(f"space after: {para.paragraph_format.space_after}")
                
            # Check for manual indentation (not from style)
            if para.paragraph_format.first_line_indent:
                issues.append(f"first line indent: {para.paragraph_format.first_line_indent}")
            if para.paragraph_format.left_indent:
                issues.append(f"left indent: {para.paragraph_format.left_indent}")
                
            # Check for center/right alignment
            if para.alignment and str(para.alignment) != "LEFT":
                issues.append(f"alignment: {para.alignment}")
                
            if issues:
                self.edge_cases.append({
                    'type': 'direct_spacing',
                    'location': f'Paragraph {i+1}',
                    'text': para.text[:100],
                    'issues': issues,
                    'severity': 'low',
                    'suggestion': 'Direct paragraph formatting applied (spacing, indent, or alignment) - review whether this should become a named paragraph style instead'
                })
    
    def detect_image_inventory(self):
        """Emit one finding per inline image so preflight can surface image counts."""
        for i, shape in enumerate(self.doc.inline_shapes):
            width_in = round(shape.width / 914400, 2) if getattr(shape, 'width', None) else None
            height_in = round(shape.height / 914400, 2) if getattr(shape, 'height', None) else None
            self.edge_cases.append({
                'type': 'image_inventory',
                'location': f'Image {i+1}',
                'text': f'Inline image {i+1}',
                'image_index': i + 1,
                'width_in': width_in,
                'height_in': height_in,
                'severity': 'low',
                'suggestion': 'Review caption / alt text / placement for this image'
            })

    def detect_observed_styles(self):
        """Inventory observed manuscript styles, but stay quiet for obvious built-ins."""
        seen = set()
        emoji_runs = []
        for i, para in enumerate(self.doc.paragraphs):
            if self._should_skip_paragraph(para):
                continue
            para_text = para.text.strip()
            para_style = getattr(getattr(para, 'style', None), 'name', None)
            para_norm = self._normalize_style_name(para_style or '')
            if para_style and para_norm not in self.BUILTIN_PARAGRAPH_STYLES:
                key = ('paragraph', para_style)
                if key not in seen:
                    seen.add(key)
                    auto_decision = 'preserve' if para_norm in self.declared_style_names else None
                    suggestion = (
                        'Declared paragraph style is already present in the manuscript — preserve it unless treatment needs to change'
                        if auto_decision == 'preserve'
                        else 'Observed paragraph style in manuscript; confirm whether it should be declared intentionally in the project spec'
                    )
                    self.edge_cases.append({
                        'type': 'observed_style',
                        'style_kind': 'paragraph',
                        'style_name': para_style,
                        'location': f'Paragraph {i+1}',
                        'text': para_text[:100] if para_text else para_style,
                        'severity': 'low',
                        'auto_decision': auto_decision,
                        'suggestion': suggestion,
                        **self._paragraph_context(i),
                    })
            for run in para.runs:
                run_text = run.text.strip()
                run_style = getattr(getattr(run, 'style', None), 'name', None)
                run_norm = self._normalize_style_name(run_style or '')
                if run_style and run_norm not in self.BUILTIN_CHARACTER_STYLES:
                    key = ('character', run_style)
                    if key not in seen:
                        seen.add(key)
                        auto_decision = 'preserve' if run_norm in self.declared_style_names else None
                        suggestion = (
                            'Declared character style is already present in the manuscript — preserve it unless treatment needs to change'
                            if auto_decision == 'preserve'
                            else 'Observed character style in manuscript; confirm whether it should be declared intentionally in the project spec'
                        )
                        self.edge_cases.append({
                            'type': 'observed_style',
                            'style_kind': 'character',
                            'style_name': run_style,
                            'location': f'Paragraph {i+1}',
                            'text': run_text[:100] if run_text else run_style,
                            'severity': 'low',
                            'auto_decision': auto_decision,
                            'suggestion': suggestion,
                            **self._paragraph_context(i),
                        })
                if self._looks_like_emoji_font(run):
                    emoji_runs.append((i + 1, run.font.name or 'emoji font', run_text[:40]))

    def detect_special_typography(self):
        """Detect likely spacing-sensitive or preformatted blocks such as ASCII art."""
        current_block = []
        current_start = None

        def flush_block():
            nonlocal current_block, current_start
            if not current_block:
                return
            block_text = '\n'.join(current_block)
            match = self._best_declared_style_match(block_text)
            auto_decision = 'preserve' if match else None
            severity = 'low' if match else 'high'
            if match:
                suggestion = (
                    f"Likely spacing-sensitive ASCII / preformatted block is already assigned to declared style "
                    f"\"{match}\" — preserve as intentional special typography. Only review further if spacing is still local formatting."
                )
            else:
                suggestion = 'Likely spacing-sensitive ASCII / preformatted block — preserve as intentional special typography'
            self.edge_cases.append({
                'type': 'special_typography',
                'subtype': 'ascii_art_candidate',
                'location': f'Paragraph {current_start}' if len(current_block) == 1 else f'Paragraphs {current_start}-{current_start + len(current_block) - 1}',
                'text': block_text[:200],
                'line_count': len(current_block),
                'declared_style_match': match,
                'severity': severity,
                'auto_decision': auto_decision,
                'suggestion': suggestion,
            })
            current_block = []
            current_start = None

        for i, para in enumerate(self.doc.paragraphs, start=1):
            if self._should_skip_paragraph(para):
                flush_block()
                continue
            text = para.text.rstrip('\n')
            if text.strip() and self._looks_like_ascii_art(text):
                if current_start is None:
                    current_start = i
                current_block.append(text)
            else:
                flush_block()

        flush_block()

    def detect_language_scripts(self):
        """Report non-Latin scripts so production can confirm language handling choices."""
        seen = set()
        for i, para in enumerate(self.doc.paragraphs):
            if self._should_skip_paragraph(para):
                continue
            text = para.text.strip()
            if not text:
                continue
            if self._looks_like_ascii_art(text) or self._is_box_drawing_heavy(text):
                continue
            script = self._detect_script(text)
            if not script:
                continue
            key = (script, text[:40])
            if key in seen:
                continue
            seen.add(key)
            self.edge_cases.append({
                'type': 'language_script',
                'location': f'Paragraph {i+1}',
                'text': text[:100],
                'script': script,
                'severity': 'medium',
                'suggestion': 'Non-Latin script detected — preserve text and verify language-specific EPUB/Typst handling'
            })

    def _looks_like_ascii_art(self, text: str) -> bool:
        stripped = text.rstrip()
        if len(stripped) < 6:
            return False
        if '  ' not in stripped and not stripped.startswith(' '):
            return False
        symbol_chars = sum(1 for ch in stripped if not ch.isalnum() and not ch.isspace())
        if symbol_chars < 2:
            return False
        if self.ASCII_DENSE_RE.search(stripped):
            return True
        unique_symbols = {ch for ch in stripped if not ch.isalnum() and not ch.isspace()}
        if len(unique_symbols) >= 2 and symbol_chars >= max(2, len(stripped) // 5):
            return True
        compact = stripped.strip()
        return any(marker in compact for marker in ('||', '//', '\\\\', '/_/', '(_', '_/')) and symbol_chars >= 2

    def _detect_script(self, text: str):
        if self.THAI_RE.search(text):
            return 'thai'
        if self.CJK_RE.search(text):
            return 'cjk'
        return None

    def _is_box_drawing_heavy(self, text: str) -> bool:
        significant = [ch for ch in text if ch.strip()]
        if not significant:
            return False
        decorative = 0
        for ch in significant:
            codepoint = ord(ch)
            if (
                0x2500 <= codepoint <= 0x257F or
                0x2580 <= codepoint <= 0x259F or
                0xFFE0 <= codepoint <= 0xFFEE or
                codepoint == 0xFF3F or
                unicodedata.name(ch, '').startswith('BOX DRAWINGS') or
                unicodedata.name(ch, '').startswith('FULLWIDTH')
            ):
                decorative += 1
        return decorative >= max(2, len(significant) // 4)

    def _normalize_style_name(self, name: str) -> str:
        return (name or '').strip().lower()

    def _find_template_guide_cutoff(self):
        start = None
        for idx, para in enumerate(self.doc.paragraphs):
            text = (para.text or '').strip()
            if start is None and text == 'Template Guide':
                start = idx
            if start is not None and text == 'Style Summary':
                return idx + 3
        return None

    def _is_template_guide_document(self) -> bool:
        texts = [p.text.strip() for p in self.doc.paragraphs[:40] if p.text.strip()]
        joined = '\n'.join(texts)
        markers = [
            'Template Guide',
            'This document is a styled Word template generated from your book specification.',
            'Custom Styles',
            'Body Text (Normal)',
            'Use \'Section Break\' to mark scene or section divisions.'
        ]
        hits = sum(1 for marker in markers if marker in joined)
        return hits >= 2

    def _should_skip_paragraph(self, para) -> bool:
        text = (para.text or '').strip()
        if self.template_guide_cutoff is not None:
            try:
                idx = self.doc.paragraphs.index(para)
                if idx <= self.template_guide_cutoff:
                    return True
            except ValueError:
                pass
        if not text:
            return False
        extra_guide_markers = (
            "This paragraph uses the 'tweet-p' style.",
            "This paragraph uses the 'metadata-p' style.",
            'metadata-c (character style):',
            'metadata-p: tweet metadata if standalone',
            '$ echo \'Hello, World!\'',
            'Shall I compare thee to a summer\'s day?',
            'In the beginning was the Word.',
        )
        return any(marker in text for marker in extra_guide_markers)

    def _best_declared_style_match(self, text: str):
        for candidate in sorted(self.declared_style_names):
            if 'ascii' in candidate:
                return candidate
        return None

    def _font_should_be_aggregated(self, font_name: str, samples: List[str]) -> bool:
        lowered = (font_name or '').lower()
        if 'emoji' in lowered or 'symbol' in lowered:
            return True
        if not samples:
            return False
        emoji_like = 0
        for sample in samples:
            for ch in sample:
                if not ch.strip():
                    continue
                if ord(ch) > 0x2600 or unicodedata.category(ch).startswith('So'):
                    emoji_like += 1
        return emoji_like >= max(1, len(''.join(samples)) // 4)

    def _looks_like_emoji_font(self, run) -> bool:
        font_name = (run.font.name or '').lower()
        if not font_name:
            return False
        return 'emoji' in font_name or 'symbol' in font_name

    def _is_from_style(self, run, attr: str) -> bool:
        """Check if formatting comes from a style (simplified)"""
        # This is a simplified check - would need more logic for full detection
        return False


class EdgeCaseReviewer:
    """Review-only interface for edge cases"""

    # Section metadata: (label, description, render_order_hint)
    SECTION_META = {
        'undeclared_custom_style': ('Undeclared Custom Styles', 'Custom styles used in manuscript but not declared in production spec'),
        'language_script': ('Language / Script', 'Non-Latin script content requiring special font or language handling'),
        'special_typography': ('Special Typography', 'Spacing-sensitive or preformatted content (ASCII art, code blocks)'),
        'observed_style': ('Observed Styles', 'Non-built-in Word styles found in the manuscript'),
        'declared_custom_style_used': ('Declared Custom Styles', 'Declared custom styles confirmed present in the manuscript'),
        'font_treatment': ('Font Treatment', 'Emoji or special font usage detected'),
        'manual_formatting': ('Manual Formatting', 'Bold, italic, or underline applied directly to runs rather than via character styles'),
        'manual_list': ('Manual Lists', 'Text formatted as lists using characters (-, *, 1.) rather than Word list styles'),
        'direct_spacing': ('Direct Paragraph Formatting', 'Paragraph-level spacing, indentation, or alignment applied directly rather than via styles'),
        'image_inventory': ('Image Inventory', 'Inline images found in the manuscript'),
        'colored_text': ('Colored Text', 'Non-black text color detected'),
        'highlighted_text': ('Highlighted Text', 'Highlighted text detected'),
        'unusual_font': ('Unusual Fonts', 'Non-standard fonts detected'),
        'manual_break': ('Manual Breaks', 'Manual section break characters detected'),
        'mixed_formatting': ('Mixed Formatting', 'Multiple fonts or sizes within a single paragraph'),
    }

    # Preferred display order (high-severity / small-count first)
    SECTION_ORDER_PRIORITY = [
        'undeclared_custom_style', 'language_script', 'special_typography',
        'colored_text', 'highlighted_text', 'unusual_font',
        'observed_style', 'declared_custom_style_used', 'font_treatment',
        'manual_formatting', 'manual_list', 'direct_spacing',
        'manual_break', 'mixed_formatting', 'image_inventory',
    ]

    def __init__(self, edge_cases: List[Dict], doc_path: str):
        self.edge_cases = edge_cases
        self.doc_path = doc_path

    # ── Helpers ────────────────────────────────────────────────────────────

    @staticmethod
    def _nice_type_label(case_type: str) -> str:
        meta = EdgeCaseReviewer.SECTION_META.get(case_type)
        return meta[0] if meta else case_type.replace('_', ' ').title()

    @staticmethod
    def _section_description(case_type: str) -> str:
        meta = EdgeCaseReviewer.SECTION_META.get(case_type)
        return meta[1] if meta else ''

    @staticmethod
    def _severity_label(severity: str) -> str:
        return {'high': 'Needs attention', 'medium': 'Review', 'low': 'Noted'}.get(severity, severity.title())

    @staticmethod
    def _severity_counts(cases: List[Dict]) -> Dict[str, int]:
        counts = {'high': 0, 'medium': 0, 'low': 0}
        for c in cases:
            sev = c.get('severity', 'low')
            counts[sev] = counts.get(sev, 0) + 1
        return counts

    @staticmethod
    def _dominant_severity(cases: List[Dict]) -> str:
        for sev in ('high', 'medium', 'low'):
            if any(c.get('severity') == sev for c in cases):
                return sev
        return 'low'

    @staticmethod
    def _esc(text) -> str:
        return html.escape(str(text)) if text else ''

    @staticmethod
    def _truncate(text: str, length: int = 80) -> str:
        text = (text or '').strip()
        return text[:length] + '…' if len(text) > length else text

    # ── Section rendering ─────────────────────────────────────────────────

    def _render_severity_badges(self, cases: List[Dict]) -> str:
        counts = self._severity_counts(cases)
        parts = []
        for sev in ('high', 'medium', 'low'):
            if counts[sev]:
                parts.append(f'<span class="badge badge-{sev}">{counts[sev]} {sev}</span>')
        return ' '.join(parts)

    def _render_bulk_table(self, section_id: str, cases: List[Dict], case_type: str) -> str:
        """Render a compact table for bulk sections (>20 items)."""
        # Build summary line
        summary = self._build_bulk_summary(cases, case_type)
        rows = []
        for c in cases:
            loc = self._esc(c.get('location', ''))
            text = self._esc(self._truncate(c.get('text', ''), 90))
            # Determine subtype column
            subtype = ''
            if case_type == 'manual_list':
                subtype = self._esc(c.get('list_type', ''))
            elif case_type == 'manual_formatting':
                subtype = self._esc(', '.join(c.get('issues', [])))
            elif case_type == 'direct_spacing':
                subtype = self._esc(', '.join(c.get('issues', []))[:60])
            elif c.get('subtype'):
                subtype = self._esc(c.get('subtype', ''))
            sev_class = c.get('severity', 'low')
            rows.append(
                f'<tr class="bulk-row" data-section="{self._esc(section_id)}">'
                f'<td class="col-loc">{loc}</td>'
                f'<td class="col-text">{text}</td>'
                f'<td class="col-sub">{subtype}</td>'
                f'<td class="col-sev"><span class="sev-dot sev-dot-{sev_class}"></span></td>'
                f'</tr>'
            )
        visible_rows = '\n'.join(rows[:10])
        hidden_rows = '\n'.join(rows[10:])
        show_more = ''
        if len(rows) > 10:
            show_more = '\n'.join(
                r.replace('class="bulk-row"', f'class="bulk-row bulk-row-hidden" style="display:none"')
                for r in rows[10:]
            )

        btn = ''
        if len(rows) > 10:
            btn = (
                f'<button class="btn-show-all" data-section="{self._esc(section_id)}" '
                f'onclick="toggleBulkRows(this, \'{self._esc(section_id)}\')" '
                f'type="button">Show all {len(rows)}</button>'
            )

        return (
            f'<p class="bulk-summary">{summary}</p>'
            f'<div class="table-wrap">'
            f'<table class="bulk-table">'
            f'<thead><tr><th>Location</th><th>Text</th><th>Detail</th><th></th></tr></thead>'
            f'<tbody>{visible_rows}\n{show_more}</tbody>'
            f'</table>'
            f'{btn}'
            f'</div>'
        )

    def _build_bulk_summary(self, cases: List[Dict], case_type: str) -> str:
        n = len(cases)
        label = self._nice_type_label(case_type).lower()
        if case_type == 'manual_list':
            by_lt = {}
            for c in cases:
                lt = c.get('list_type', 'unknown')
                by_lt[lt] = by_lt.get(lt, 0) + 1
            parts = ', '.join(f'{v} {k}' for k, v in sorted(by_lt.items(), key=lambda x: -x[1]))
            return f'{n} {label}: {parts}'
        elif case_type == 'manual_formatting':
            issue_counts = {}
            for c in cases:
                for iss in c.get('issues', []):
                    issue_counts[iss] = issue_counts.get(iss, 0) + 1
            parts = ', '.join(f'{v} {k}' for k, v in sorted(issue_counts.items(), key=lambda x: -x[1]))
            return f'{n} {label} runs: {parts}'
        elif case_type == 'direct_spacing':
            return f'{n} paragraphs with direct spacing, indentation, or alignment'
        return f'{n} {label} items'

    def _render_card(self, case: Dict) -> str:
        """Render a single finding as a compact card."""
        sev = case.get('severity', 'low')
        parts = []
        parts.append(f'<article class="finding severity-{sev}">')
        parts.append(
            f'<div class="finding-top">'
            f'<div class="location">{self._esc(case.get("location", ""))}</div>'
            f'<div class="severity-note">{self._severity_label(sev)}</div>'
            f'</div>'
        )
        parts.append(f'<div class="text-sample">{self._esc(case.get("text", ""))}</div>')

        if case.get('script'):
            parts.append(f'<span class="script-tag">{self._esc(case["script"])}</span>')
        if case.get('subtype'):
            parts.append(f'<div class="detail"><strong>Subtype:</strong> {self._esc(case["subtype"])}</div>')
        if case.get('declared_style_match'):
            parts.append(f'<div class="detail"><strong>Declared style match:</strong> {self._esc(case["declared_style_match"])}</div>')
        if case.get('style_name'):
            kind = self._esc(case.get('style_kind', 'style'))
            parts.append(f'<div class="detail"><strong>Observed {kind} style:</strong> {self._esc(case["style_name"])}</div>')
        if case.get('context_before'):
            ctx = ' / '.join(self._esc(x) for x in case['context_before'])
            parts.append(f'<div class="detail"><strong>Before:</strong> {ctx}</div>')
        if case.get('context_after'):
            ctx = ' / '.join(self._esc(x) for x in case['context_after'])
            parts.append(f'<div class="detail"><strong>After:</strong> {ctx}</div>')
        if case.get('issues'):
            parts.append(f'<div class="detail"><strong>Issues:</strong> {self._esc(", ".join(case["issues"]))}</div>')
        if case.get('font'):
            parts.append(f'<div class="detail"><strong>Font:</strong> {self._esc(case["font"])}</div>')
        if case.get('color'):
            parts.append(f'<div class="detail"><strong>Color:</strong> {self._esc(case["color"])}</div>')
        if case.get('suggestion'):
            parts.append(f'<div class="suggestion">{self._esc(case["suggestion"])}</div>')
        parts.append('</article>')
        return '\n'.join(parts)

    def _render_card_section(self, cases: List[Dict]) -> str:
        """Render ≤20 items as compact cards inside a finding-list container."""
        inner = '\n'.join(self._render_card(c) for c in cases)
        return f'<div class="finding-list">{inner}</div>'

    def _render_image_grid(self, cases: List[Dict]) -> str:
        """Render image inventory as a compact grid."""
        if not cases:
            return (
                '<div class="empty-note">'
                '<p>No inline images detected.</p>'
                '<p class="detail">If you expected images, check whether they are floating/anchored shapes rather than inline Word images.</p>'
                '</div>'
            )
        cells = []
        for c in cases:
            idx = c.get('image_index', '?')
            w = c.get('width_in')
            h = c.get('height_in')
            dims = f'{w}″ × {h}″' if w and h else 'unknown size'
            cells.append(
                f'<div class="img-cell">'
                f'<div class="img-cell-icon">🖼</div>'
                f'<div class="img-cell-label">Image {self._esc(str(idx))}</div>'
                f'<div class="img-cell-dims">{self._esc(dims)}</div>'
                f'</div>'
            )
        joined_cells = '\n'.join(cells)
        return f'<div class="img-grid">{joined_cells}</div>'

    def _render_auto_preserved(self, cases: List[Dict]) -> str:
        """Render the auto-preserved section as grouped compact tables."""
        if not cases:
            return (
                '<div class="empty-note">'
                'No auto-preserved formatting items were detected on this run.'
                '</div>'
            )

        # Group by issue type
        groups: Dict[str, List[Dict]] = {}
        for c in cases:
            issues = c.get('issues', [])
            key = ', '.join(issues) if issues else c.get('type', 'other')
            groups.setdefault(key, []).append(c)

        # Sort groups: largest first
        sorted_groups = sorted(groups.items(), key=lambda x: -len(x[1]))

        parts = []
        parts.append(
            '<p class="bulk-summary" style="margin-bottom:0.75rem">'
            f'{len(cases)} items detected but treated as intentional and preserved.</p>'
        )

        for group_label, group_cases in sorted_groups:
            nice_label = self._esc(group_label.replace('_', ' '))
            group_id = f'ap-{hash(group_label) % 99999}'
            parts.append(
                f'<div class="ap-group" style="margin-bottom:0.75rem">'
                f'<div class="ap-group-header" style="font-size:0.78rem;font-family:var(--sans);'
                f'color:var(--text-secondary);margin-bottom:0.35rem;">'
                f'{nice_label} <span style="color:var(--text-muted)">({len(group_cases)})</span></div>'
            )
            rows = []
            has_detail = any(c.get('style_name') or c.get('declared_style_match') for c in group_cases)
            for c in group_cases:
                loc = self._esc(c.get('location', ''))
                text = self._esc(self._truncate(c.get('text', ''), 80))
                detail_parts = []
                if c.get('declared_style_match'):
                    detail_parts.append(self._esc(c['declared_style_match']))
                elif c.get('style_name'):
                    detail_parts.append(self._esc(c['style_name']))
                detail_td = f'<td class="col-sub">{", ".join(detail_parts)}</td>' if has_detail else ''
                rows.append(
                    f'<tr class="bulk-row" data-section="{self._esc(group_id)}">'
                    f'<td class="col-loc">{loc}</td>'
                    f'<td class="col-text">{text}</td>'
                    f'{detail_td}'
                    f'</tr>'
                )
            # Show first 5 rows, hide rest
            visible = '\n'.join(rows[:5])
            hidden = ''
            if len(rows) > 5:
                hidden = '\n'.join(
                    r.replace('class="bulk-row"',
                              f'class="bulk-row bulk-row-hidden" style="display:none"')
                    for r in rows[5:]
                )
            btn = ''
            if len(rows) > 5:
                btn = (
                    f'<button class="btn-show-all" data-section="{self._esc(group_id)}" '
                    f'onclick="toggleBulkRows(this, \'{self._esc(group_id)}\')" '
                    f'type="button">Show all {len(rows)}</button>'
                )
            detail_th = '<th>Detail</th>' if has_detail else ''
            parts.append(
                f'<div class="table-wrap">'
                f'<table class="bulk-table">'
                f'<thead><tr><th>Location</th><th>Text</th>{detail_th}</tr></thead>'
                f'<tbody>{visible}\n{hidden}</tbody>'
                f'</table>'
                f'{btn}'
                f'</div>'
                f'</div>'
            )

        return '\n'.join(parts)

    # ── Main report generation ────────────────────────────────────────────

    def generate_html_report(self, output_path: str):
        """Generate an HTML report for review"""
        auto_preserve = [c for c in self.edge_cases if c.get('auto_decision') == 'preserve']
        actionable = [c for c in self.edge_cases if c.get('auto_decision') != 'preserve']

        report_generated_at = datetime.now().astimezone().strftime('%Y-%m-%d %H:%M %Z')
        report_title = Path(self.doc_path).name

        # Partition actionable findings by type
        by_type: Dict[str, List[Dict]] = {}
        for c in actionable:
            by_type.setdefault(c['type'], []).append(c)
        # Ensure image_inventory section always exists
        if 'image_inventory' not in by_type:
            by_type['image_inventory'] = []

        # Order sections: priority list first, then any remaining
        seen = set()
        section_order = []
        for t in self.SECTION_ORDER_PRIORITY:
            if t in by_type:
                section_order.append(t)
                seen.add(t)
        for t in by_type:
            if t not in seen:
                section_order.append(t)

        # Severity totals for overview
        sev_totals = self._severity_counts(actionable)

        # Build stacked bar data
        bar_segments = []
        total_actionable = len(actionable)
        for t in section_order:
            cases = by_type[t]
            if cases:
                bar_segments.append((t, len(cases), self._dominant_severity(cases)))

        # ── Build HTML ─────────────────────────────────────────────────
        sections_html = []
        nav_pills = []

        for case_type in section_order:
            cases = by_type[case_type]
            count = len(cases)
            label = self._nice_type_label(case_type)
            desc = self._section_description(case_type)
            sid = f'sec-{case_type}'
            sev_badges = self._render_severity_badges(cases)
            dom_sev = self._dominant_severity(cases)

            # Decide if section starts expanded
            has_high = any(c.get('severity') == 'high' for c in cases)
            starts_open = (count <= 3 and count > 0) or has_high
            open_attr = ' open' if starts_open else ''

            # Build inner content
            if case_type == 'image_inventory':
                inner = self._render_image_grid(cases)
            elif count > 20 and case_type in ('manual_formatting', 'manual_list', 'direct_spacing', 'mixed_formatting'):
                inner = self._render_bulk_table(sid, cases, case_type)
            elif count > 0:
                inner = self._render_card_section(cases)
            else:
                inner = '<div class="empty-note">No items detected.</div>'

            # Nav pill
            nav_pills.append(
                f'<a href="#{sid}" class="nav-pill" data-section="{self._esc(sid)}">'
                f'{self._esc(label)} <span class="nav-count">{count}</span></a>'
            )

            sections_html.append(
                f'<section id="{self._esc(sid)}" class="report-section">'
                f'<details{open_attr}>'
                f'<summary class="section-summary">'
                f'<div class="section-summary-left">'
                f'<span class="section-chevron"></span>'
                f'<span class="section-label-inline">{self._esc(label)}</span>'
                f'<span class="badge badge-dim">{count}</span>'
                f'{sev_badges}'
                f'</div>'
                f'<div class="section-summary-desc">{self._esc(desc)}</div>'
                f'</summary>'
                f'<div class="section-body">{inner}</div>'
                f'</details>'
                f'</section>'
            )

        # Auto-preserved section
        ap_count = len(auto_preserve)
        ap_open = ' open' if ap_count <= 3 and ap_count > 0 else ''
        ap_inner = self._render_auto_preserved(auto_preserve)
        nav_pills.append(
            f'<a href="#sec-auto-preserved" class="nav-pill" data-section="sec-auto-preserved">'
            f'Auto-preserved <span class="nav-count">{ap_count}</span></a>'
        )
        sections_html.append(
            f'<section id="sec-auto-preserved" class="report-section">'
            f'<details{ap_open}>'
            f'<summary class="section-summary">'
            f'<div class="section-summary-left">'
            f'<span class="section-chevron"></span>'
            f'<span class="section-label-inline">Auto-Preserved</span>'
            f'<span class="badge badge-dim">{ap_count}</span>'
            f'</div>'
            f'<div class="section-summary-desc">Items detected but treated as intentional — no action needed</div>'
            f'</summary>'
            f'<div class="section-body">{ap_inner}</div>'
            f'</details>'
            f'</section>'
        )

        # Stacked bar HTML
        # Per-section color palette (muted, distinguishable hues)
        section_hues = [
            '#6366f1', '#8b5cf6', '#ec4899', '#f43f5e',
            '#f97316', '#eab308', '#22c55e', '#14b8a6',
            '#06b6d4', '#3b82f6', '#a855f7', '#ef4444',
        ]

        bar_html = ''
        if total_actionable > 0:
            max_n = max(n for (_, n, _) in bar_segments) if bar_segments else 1
            bar_rows = []
            for i, (t, n, sev) in enumerate(bar_segments):
                pct = (n / max_n) * 100
                lbl = self._nice_type_label(t)
                color = section_hues[i % len(section_hues)]
                bar_rows.append(
                    f'<div class="section-bar-row">'
                    f'<span class="section-bar-label">{self._esc(lbl)}</span>'
                    f'<span class="section-bar-track">'
                    f'<span class="section-bar-fill" style="width:{pct:.1f}%;background:{color}"></span>'
                    f'</span>'
                    f'<span class="section-bar-count">{n}</span>'
                    f'</div>'
                )
            bar_html = f'<div class="section-bars">{"" .join(bar_rows)}</div>'

        # ── Final assembly ─────────────────────────────────────────────
        html_out = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Preflight — {self._esc(report_title)}</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=IBM+Plex+Sans:wght@400;500;600&family=IBM+Plex+Serif:ital,wght@0,400;0,500;0,600;1,400&display=swap" rel="stylesheet">
<script>
(function(){{ try {{ var t=JSON.parse(localStorage.getItem('preflight-theme-v1')); if(t&&t.dark) document.documentElement.classList.add('dark'); }} catch(e){{}} }})();
</script>
<style>
/* ── Reset + CSS Variables ───────────────────────────────── */
* {{ margin:0; padding:0; box-sizing:border-box; }}

:root {{
  --bg:#f8f9fa; --surface:#ffffff; --border:#e0e2e6;
  --text:#111827; --text-secondary:#6b7280; --text-muted:#9ca3af;
  --accent:#2563eb; --green:#16a34a; --red:#dc2626; --yellow:#ca8a04; --orange:#ea580c;
  --badge-green-bg:rgba(22,163,74,0.1); --badge-red-bg:rgba(220,38,38,0.1);
  --badge-yellow-bg:rgba(202,138,4,0.1);
  --badge-dim-bg:#f1f3f5; --progress-bg:#f1f3f5;
  --noise-opacity:0.035; --grid-color:rgba(160,168,180,0.6);
  --gradient-wash:linear-gradient(to bottom,rgba(230,233,240,0.3) 0%,transparent 60%);
  --theme-bg:rgba(255,255,255,0.92);
  --sans:'IBM Plex Sans',Arial,sans-serif;
  --serif:'IBM Plex Serif',Georgia,serif;
  --mono:'SFMono-Regular',Menlo,Consolas,'Liberation Mono',monospace;
}}

html.dark {{
  --bg:#0f1115; --surface:#1a1d24; --border:#2a2d35;
  --text:#e4e5e7; --text-secondary:#9ca0a8; --text-muted:#6b7074;
  --accent:#5b9aff; --green:#34d399; --red:#f87171; --yellow:#fbbf24; --orange:#fb923c;
  --badge-green-bg:rgba(52,211,153,0.12); --badge-red-bg:rgba(248,113,113,0.12);
  --badge-yellow-bg:rgba(251,191,36,0.12);
  --badge-dim-bg:#23262e; --progress-bg:#23262e;
  --noise-opacity:0.012; --grid-color:rgba(90,96,110,0.7);
  --gradient-wash:linear-gradient(to bottom,rgba(20,22,28,0.4) 0%,transparent 60%);
  --theme-bg:rgba(26,29,36,0.92);
}}

/* ── Base ────────────────────────────────────────────────── */
body {{
  background:var(--bg); color:var(--text);
  font-family:var(--serif); line-height:1.55;
  min-height:100vh; position:relative;
  transition:background .3s,color .3s;
}}
body, body * {{ font-weight:400; }}

body::before {{
  content:''; position:fixed; inset:0; z-index:0; pointer-events:none;
  background-image:url("data:image/svg+xml,%3Csvg viewBox='0 0 256 256' xmlns='http://www.w3.org/2000/svg'%3E%3Cfilter id='n'%3E%3CfeTurbulence type='fractalNoise' baseFrequency='0.9' numOctaves='4' stitchTiles='stitch'/%3E%3C/filter%3E%3Crect width='100%25' height='100%25' filter='url(%23n)'/%3E%3C/svg%3E");
  background-size:180px; opacity:var(--noise-opacity);
}}
body::after {{
  content:''; position:fixed; top:0; left:0; right:0; height:50vh;
  z-index:0; pointer-events:none; background:var(--gradient-wash);
}}

.grid-lines {{ position:fixed; inset:0; z-index:0; pointer-events:none; display:flex; justify-content:center; }}
.grid-lines-inner {{ width:1120px; max-width:100vw; display:flex; justify-content:space-between; }}
.grid-line {{ width:1px; background:repeating-linear-gradient(to bottom,var(--grid-color) 0px,var(--grid-color) 4px,transparent 4px,transparent 10px); }}

.content {{ position:relative; z-index:1; max-width:1100px; margin:0 auto; padding:2.5rem 2rem 4rem; }}

/* ── Typography ──────────────────────────────────────────── */
h1,h2,h3 {{ margin:0; font-weight:400; }}
p,li {{ font-size:0.95rem; }}

.eyebrow {{
  font-size:0.7rem; color:var(--text-muted); text-transform:uppercase;
  letter-spacing:0.08em; font-family:var(--sans); margin-bottom:0.5rem;
}}

/* ── Header ──────────────────────────────────────────────── */
.header {{
  padding-bottom:1.25rem; border-bottom:1px solid var(--border); margin-bottom:1.25rem;
}}
.header h1 {{ font-size:1.5rem; letter-spacing:-0.03em; }}
.meta {{ margin-top:0.65rem; color:var(--text-secondary); font-size:0.9rem; }}

/* ── Dark mode toggle ────────────────────────────────────── */
.dark-toggle {{
  position:fixed; top:1rem; right:2rem; z-index:200;
  width:28px; height:28px; border-radius:50%; border:1px solid var(--border);
  background:var(--theme-bg); backdrop-filter:blur(12px); -webkit-backdrop-filter:blur(12px);
  color:var(--text-secondary); cursor:pointer;
  display:flex; align-items:center; justify-content:center;
  font-size:13px; transition:all .2s;
  box-shadow:0 2px 8px rgba(0,0,0,0.08);
}}
.dark-toggle:hover {{ color:var(--text); border-color:var(--accent); }}

/* ── Sticky nav ──────────────────────────────────────────── */
.sticky-nav {{
  position:sticky; top:0; z-index:100;
  background:var(--theme-bg); backdrop-filter:blur(12px); -webkit-backdrop-filter:blur(12px);
  border-bottom:1px solid var(--border);
  margin:0 -2rem 0; padding:0.5rem 2rem;
  display:flex; gap:0.35rem; overflow-x:auto;
  opacity:0; pointer-events:none; transition:opacity .25s, max-height .25s;
  max-height:0; overflow:hidden;
  -ms-overflow-style:none; scrollbar-width:none;
}}
.sticky-nav::-webkit-scrollbar {{ display:none; }}
.sticky-nav.visible {{ opacity:1; pointer-events:auto; max-height:3rem; overflow-x:auto; overflow-y:hidden; }}

.nav-pill {{
  flex-shrink:0; padding:0.2rem 0.6rem; border-radius:9999px;
  font-size:0.68rem; font-family:var(--sans); color:var(--text-secondary);
  text-decoration:none; white-space:nowrap;
  border:1px solid transparent; transition:all .15s;
}}
.nav-pill:hover {{ border-color:var(--border); background:var(--surface); }}
.nav-pill.active {{ border-color:var(--accent); color:var(--accent); background:rgba(37,99,235,0.06); }}
html.dark .nav-pill.active {{ background:rgba(91,154,255,0.1); }}
.nav-count {{
  display:inline-block; min-width:1.1em; text-align:center;
  font-size:0.62rem; color:var(--text-muted);
}}

/* ── Overview card ────────────────────────────────────────── */
.overview {{
  background:var(--surface); border:1px solid var(--border); border-radius:6px;
  padding:1.25rem 1.5rem; margin-bottom:1.5rem;
}}
.overview-label {{
  font-size:0.7rem; color:var(--text-muted); text-transform:uppercase;
  letter-spacing:0.08em; font-family:var(--sans); margin-bottom:0.6rem;
}}
.overview h2 {{ font-size:1.05rem; margin-bottom:0.25rem; }}
.overview-copy {{ color:var(--text-secondary); font-size:0.88rem; margin-bottom:0.85rem; }}
.summary-grid {{ display:flex; flex-wrap:wrap; gap:0.4rem; margin-bottom:1rem; }}

/* Section breakdown bars */
.section-bars {{ display:flex; flex-direction:column; gap:0.35rem; }}
.section-bar-row {{
  display:flex; align-items:center; gap:0.5rem;
  font-size:0.72rem; font-family:var(--sans);
}}
.section-bar-label {{
  width:160px; flex-shrink:0; text-align:right;
  color:var(--text-secondary); white-space:nowrap;
  overflow:hidden; text-overflow:ellipsis;
}}
.section-bar-track {{
  flex:1; height:6px; border-radius:3px;
  background:var(--progress-bg); overflow:hidden;
}}
.section-bar-fill {{
  height:100%; border-radius:3px; transition:width .3s;
}}
.section-bar-count {{
  width:32px; flex-shrink:0; text-align:right;
  color:var(--text-muted); font-size:0.68rem;
  font-variant-numeric:tabular-nums;
}}
@media (max-width:800px) {{
  .section-bar-label {{ width:100px; font-size:0.65rem; }}
}}

/* ── Badges ──────────────────────────────────────────────── */
.badge {{
  display:inline-flex; align-items:center; gap:0.3rem;
  padding:0.18rem 0.55rem; border-radius:9999px;
  font-size:0.7rem; font-family:var(--sans);
  background:var(--badge-dim-bg); color:var(--text-secondary);
}}
.badge-high {{ background:var(--badge-red-bg); color:var(--red); }}
.badge-medium {{ background:var(--badge-yellow-bg); color:var(--yellow); }}
.badge-low {{ background:var(--badge-green-bg); color:var(--green); }}

/* ── Section (collapsible details) ───────────────────────── */
.report-section {{
  margin-bottom:0.75rem;
}}
.report-section details {{
  background:var(--surface); border:1px solid var(--border); border-radius:6px;
  overflow:hidden;
}}
.section-summary {{
  list-style:none; cursor:pointer; padding:0.85rem 1.25rem;
  display:flex; flex-wrap:wrap; align-items:center;
  gap:0.4rem 0.6rem; user-select:none;
  transition:background .15s;
}}
.section-summary:hover {{ background:var(--progress-bg); }}
.section-summary::-webkit-details-marker {{ display:none; }}

.section-chevron {{
  display:inline-block; width:12px; height:12px; flex-shrink:0;
  transition:transform .2s;
  background:none; position:relative;
}}
.section-chevron::after {{
  content:''; position:absolute; top:3px; left:2px;
  width:6px; height:6px; border-right:1.5px solid var(--text-muted);
  border-bottom:1.5px solid var(--text-muted); transform:rotate(-45deg);
  transition:transform .2s;
}}
details[open] > .section-summary .section-chevron::after {{
  transform:rotate(45deg); top:2px;
}}

.section-summary-left {{
  display:flex; align-items:center; gap:0.5rem; flex-shrink:0;
}}
.section-label-inline {{
  font-size:0.88rem; font-family:var(--serif); color:var(--text);
}}
.section-summary-desc {{
  font-size:0.78rem; color:var(--text-muted); font-family:var(--sans);
  margin-left:auto;
}}

.section-body {{
  padding:0 1.25rem 1.25rem;
}}

/* ── Finding cards ───────────────────────────────────────── */
.finding-list {{
  border:1px solid var(--border); border-radius:4px; overflow:hidden;
}}
.finding {{
  padding:0.75rem 0.9rem; border-top:1px solid var(--border);
}}
.finding:first-child {{ border-top:none; }}
.finding-top {{
  display:flex; justify-content:space-between; align-items:baseline;
  gap:0.75rem; margin-bottom:0.35rem;
}}
.location {{ font-size:0.8rem; color:var(--text-secondary); font-family:var(--sans); }}
.severity-note {{
  font-size:0.68rem; color:var(--text-muted); text-transform:uppercase;
  letter-spacing:0.05em; font-family:var(--sans);
}}
.severity-high .severity-note {{ color:var(--red); }}
.severity-medium .severity-note {{ color:var(--orange); }}
.severity-low .severity-note {{ color:var(--green); }}

.text-sample {{
  background:var(--bg); border:1px solid var(--border); border-radius:3px;
  padding:0.5rem 0.65rem; margin:0.4rem 0;
  font-family:var(--mono); font-size:0.8rem;
  white-space:pre-wrap; overflow-wrap:anywhere;
  max-height:4.5em; overflow-y:auto;
}}
.detail {{ font-size:0.82rem; color:var(--text-secondary); margin:0.25rem 0; }}
.detail strong {{ color:var(--text); font-weight:500; }}
.detail-inline {{ font-size:0.8rem; color:var(--text-muted); }}
.suggestion {{ font-size:0.82rem; color:var(--text-muted); margin-top:0.35rem; font-style:italic; }}
.script-tag {{
  display:inline-block; padding:0.15rem 0.5rem; border-radius:999px;
  background:rgba(37,99,235,0.08); color:var(--accent);
  font-size:0.7rem; font-family:var(--sans); margin-bottom:0.3rem;
}}

/* ── Bulk table ──────────────────────────────────────────── */
.bulk-summary {{ font-size:0.88rem; color:var(--text-secondary); margin-bottom:0.75rem; }}
.table-wrap {{ overflow-x:auto; }}
.bulk-table {{
  width:100%; border-collapse:collapse; font-size:0.8rem;
  font-family:var(--sans);
}}
.bulk-table th {{
  text-align:left; font-size:0.65rem; text-transform:uppercase;
  letter-spacing:0.06em; color:var(--text-muted); padding:0.45rem 0.6rem;
  border-bottom:1px solid var(--border); font-weight:400;
}}
.bulk-table td {{
  padding:0.35rem 0.6rem; border-bottom:1px solid var(--border);
  color:var(--text-secondary); vertical-align:top;
}}
.col-loc {{ white-space:nowrap; width:100px; color:var(--text-muted); font-size:0.75rem; }}
.col-text {{ font-family:var(--mono); font-size:0.75rem; max-width:500px; overflow-wrap:anywhere; }}
.col-sub {{ font-size:0.72rem; color:var(--text-muted); max-width:180px; overflow-wrap:anywhere; }}
.col-sev {{ width:20px; text-align:center; }}
.sev-dot {{ display:inline-block; width:6px; height:6px; border-radius:50%; }}
.sev-dot-high {{ background:var(--red); }}
.sev-dot-medium {{ background:var(--yellow); }}
.sev-dot-low {{ background:var(--green); }}

.btn-show-all {{
  margin-top:0.5rem; padding:0.35rem 0.9rem; border-radius:4px;
  border:1px solid var(--border); background:var(--surface); color:var(--text-secondary);
  font-size:0.78rem; font-family:var(--sans); cursor:pointer; transition:all .15s;
}}
.btn-show-all:hover {{ border-color:var(--accent); color:var(--accent); }}

/* ── Image grid ──────────────────────────────────────────── */
.img-grid {{
  display:grid; grid-template-columns:repeat(auto-fill,minmax(140px,1fr));
  gap:0.5rem;
}}
.img-cell {{
  background:var(--bg); border:1px solid var(--border); border-radius:4px;
  padding:0.65rem; text-align:center;
}}
.img-cell-icon {{ font-size:1.4rem; margin-bottom:0.15rem; }}
.img-cell-label {{ font-size:0.78rem; font-family:var(--sans); color:var(--text); }}
.img-cell-dims {{ font-size:0.7rem; color:var(--text-muted); font-family:var(--sans); }}

.empty-note {{
  background:var(--surface); border:1px solid var(--border); border-radius:4px;
  padding:1rem; color:var(--text-secondary); font-size:0.88rem;
}}
.empty-note .detail {{ font-size:0.82rem; margin-top:0.35rem; }}

/* ── Responsive ──────────────────────────────────────────── */
@media (max-width:800px) {{
  .content {{ padding:1.5rem 1rem 3rem; }}
  .sticky-nav {{ margin:0 -1rem 1.5rem; padding:0.5rem 1rem; }}
  .finding-top {{ flex-direction:column; align-items:flex-start; gap:0.2rem; }}
  .section-summary {{ padding:0.7rem 1rem; }}
  .section-body {{ padding:0 1rem 1rem; }}
  .section-summary-desc {{ margin-left:0; margin-top:0.15rem; }}
  .dark-toggle {{ top:0.5rem; right:0.75rem; }}
  .img-grid {{ grid-template-columns:repeat(auto-fill,minmax(110px,1fr)); }}
}}

/* ── Print ───────────────────────────────────────────────── */
@media print {{
  .dark-toggle, .sticky-nav, .grid-lines {{ display:none !important; }}
  body::before, body::after {{ display:none !important; }}
  details {{ display:block !important; }}
  details > summary {{ display:none !important; }}
  details > .section-body {{ display:block !important; }}
  .report-section details {{ border:1px solid #ccc; page-break-inside:avoid; }}
  .bulk-row-hidden {{ display:table-row !important; }}
  .btn-show-all {{ display:none !important; }}
  .content {{ max-width:100%; padding:1rem; }}
  .section-body {{ padding:0.5rem 1rem 1rem !important; }}
  /* Print needs a header for each section */
  .report-section::before {{
    content:attr(data-print-label);
    display:block; font-size:0.8rem; font-weight:500;
    text-transform:uppercase; letter-spacing:0.06em;
    color:#666; margin-bottom:0.5rem; padding-top:0.75rem;
    border-top:1px solid #ccc;
  }}
}}
</style>
</head>
<body>
<div class="grid-lines"><div class="grid-lines-inner"><div class="grid-line"></div><div class="grid-line"></div></div></div>

<button class="dark-toggle" id="darkToggle" type="button" aria-label="Toggle dark mode">☀</button>

<main class="content">
  <header class="header">
    <div class="eyebrow">Production · Typesetting</div>
    <h1>{self._esc(report_title)}</h1>
    <div class="meta">Manuscript preflight check before generating EPUB or PDF.</div>
    <div class="meta">Generated {report_generated_at}</div>
  </header>

  <nav class="sticky-nav" id="stickyNav">
    {''.join(nav_pills)}
  </nav>

  <section class="overview">
    <div class="overview-label">Overview</div>
    <h2>Preflight summary</h2>
    <p class="overview-copy">Review findings, make fixes in Word if needed, then rerun inspection.</p>
    <div class="summary-grid">
      <span class="badge">{len(self.edge_cases)} total</span>
      <span class="badge">{len(actionable)} actionable</span>
      <span class="badge">{ap_count} auto-preserved</span>
      <span class="badge badge-high">{sev_totals['high']} high</span>
      <span class="badge badge-medium">{sev_totals['medium']} medium</span>
      <span class="badge badge-low">{sev_totals['low']} low</span>
    </div>
    {bar_html}
  </section>

  {''.join(sections_html)}
</main>

<script>
/* ── Dark mode ───────────────────────────────────────────── */
(function() {{
  var KEY = 'preflight-theme-v1';
  var btn = document.getElementById('darkToggle');
  function isDark() {{ return document.documentElement.classList.contains('dark'); }}
  function updateIcon() {{ btn.textContent = isDark() ? '☾' : '☀'; }}
  updateIcon();
  btn.addEventListener('click', function() {{
    document.documentElement.classList.toggle('dark');
    try {{ localStorage.setItem(KEY, JSON.stringify({{ dark: isDark() }})); }} catch(e) {{}}
    updateIcon();
  }});
}})();

/* ── Sticky nav visibility ───────────────────────────────── */
(function() {{
  var nav = document.getElementById('stickyNav');
  var header = document.querySelector('.header');
  if (!nav || !header) return;
  var io = new IntersectionObserver(function(entries) {{
    nav.classList.toggle('visible', !entries[0].isIntersecting);
  }}, {{ threshold: 0 }});
  io.observe(header);
}})();

/* ── Active nav pill on scroll ───────────────────────────── */
(function() {{
  var pills = document.querySelectorAll('.nav-pill');
  var sections = document.querySelectorAll('.report-section');
  if (!pills.length || !sections.length) return;
  var sectionMap = {{}};
  pills.forEach(function(p) {{ sectionMap[p.getAttribute('data-section')] = p; }});

  function update() {{
    var current = null;
    var scrollY = window.scrollY || window.pageYOffset;
    sections.forEach(function(s) {{
      if (s.getBoundingClientRect().top <= 120) current = s.id;
    }});
    pills.forEach(function(p) {{ p.classList.remove('active'); }});
    if (current && sectionMap[current]) sectionMap[current].classList.add('active');
  }}
  var ticking = false;
  window.addEventListener('scroll', function() {{
    if (!ticking) {{ requestAnimationFrame(function() {{ update(); ticking = false; }}); ticking = true; }}
  }});
  update();
}})();

/* ── Smooth scroll for nav pills ─────────────────────────── */
(function() {{
  document.querySelectorAll('.nav-pill').forEach(function(a) {{
    a.addEventListener('click', function(e) {{
      e.preventDefault();
      var target = document.querySelector(a.getAttribute('href'));
      if (!target) return;
      var details = target.querySelector('details');
      if (details && !details.open) details.open = true;
      target.scrollIntoView({{ behavior: 'smooth', block: 'start' }});
    }});
  }});
}})();

/* ── Bulk table show-all toggle ──────────────────────────── */
function toggleBulkRows(btn, sectionId) {{
  var rows = document.querySelectorAll('.bulk-row-hidden[data-section="' + sectionId + '"]');
  var showing = btn.getAttribute('data-expanded') === '1';
  rows.forEach(function(r) {{ r.style.display = showing ? 'none' : ''; }});
  btn.setAttribute('data-expanded', showing ? '0' : '1');
  // Count visible rows (non-hidden siblings) to get correct total
  var table = btn.closest('.table-wrap').querySelector('table');
  var allRows = table ? table.querySelectorAll('tbody tr') : [];
  var total = allRows.length;
  btn.textContent = showing ? 'Show all ' + total : 'Show fewer';
}}

/* ── Hash navigation: auto-open targeted section ─────────── */
(function() {{
  function openHash() {{
    var hash = window.location.hash;
    if (!hash) return;
    var target = document.querySelector(hash);
    if (!target) return;
    var details = target.querySelector('details');
    if (details && !details.open) details.open = true;
    setTimeout(function() {{
      target.scrollIntoView({{ behavior: 'smooth', block: 'start' }});
    }}, 100);
  }}
  openHash();
  window.addEventListener('hashchange', openHash);
}})();
</script>
</body>
</html>"""

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_out)


def main():
    parser = argparse.ArgumentParser(description='Detect edge cases in Word documents')
    parser.add_argument('input_file', help='Input Word document')
    parser.add_argument('-o', '--output', default='edge_case_review.html',
                       help='Output HTML report (default: edge_case_review.html)')
    parser.add_argument('--json', action='store_true',
                       help='Also output raw JSON data')
    parser.add_argument('--declared-styles', default='',
                       help='Optional JSON file of declared custom styles from transmittal/spec')
    
    args = parser.parse_args()
    
    declared_styles = []
    if args.declared_styles:
        declared_path = Path(args.declared_styles)
        if declared_path.exists():
            declared_styles = json.loads(declared_path.read_text(encoding='utf-8'))
    
    print(f"Analyzing {args.input_file}...")
    detector = EdgeCaseDetector(args.input_file, declared_styles=declared_styles)
    edge_cases = detector.detect_all()
    
    print(f"Found {len(edge_cases)} edge cases")
    
    # Generate HTML review interface
    reviewer = EdgeCaseReviewer(edge_cases, args.input_file)
    reviewer.generate_html_report(args.output)
    print(f"HTML report generated: {args.output}")
    
    # Optionally save JSON
    if args.json:
        json_output = args.output.replace('.html', '.json')
        with open(json_output, 'w') as f:
            json.dump(edge_cases, f, indent=2)
        print(f"JSON data saved: {json_output}")


if __name__ == '__main__':
    main()